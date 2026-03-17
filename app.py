# Date Created DD/MM/YYYY : 17/03/2026
# Author : Samar Kumar [Sam Darkwood]
# Gmail : samarkumarchoudhary823@gmail.com 

import streamlit as st 
from groq import Groq
from docxtpl import DocxTemplate
import datetime 
import os 
import io

GROQ_API_KEY = "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"
client = Groq(api_key=GROQ_API_KEY)

# --- PAGE UI SETUP --- 
st.set_page_config(page_title="Valedion AI", page_icon="", layout="wide")

st.markdown(
    """
    <style> 
    .main { 
        background-color: #0e1117;
        color: #ffffff;
    }

    .stButton>button {
        width: 100%;
        border-radius: 5px;
        height: 3em;
        background-color: #ff4b4b;
        color: white;
        font-weight: bold;
    }

    .stTextArea>div>div>textarea {
    background-color: #161b22;
    color: #00ff41;
    font-family: 'Courier New', Courier, monospace;
    }
    </style>
""", 
unsafe_allow_html=True
)

st.title("Valedion AI Powered Report generator [Privacy First]")
st.write("Professional Audit Reports from Logs in Seconds.") 

# --- SIDEBAR SETTINGS --- 
with st.sidebar:
    st.header("Report Details") 
    client_name = st.text_input("Client Company Name", "Acme Corp")
    analyst_name = st.text_input("Lead Analyst", "Cyber Security Team")
    report_type = st.selectbox("Audit Type", ["External Pentest", "Internal Vulnerability Scan", "Web Applcation  Audit"]) 
    model_choice = st.selectbox("AI Engine", ["llama-3.3-70b-versatile", "llama-3.1-8b-instant"])

    st.markdown("---")
    st.info(" **Demo Tip: ** Paste an Nmap scan or a Nessus export into the main box to see the AI magic.") 


# --- Main Interface ---
col1, col2 = st.columns([2, 1])

with col1:
    raw_logs = st.text_area("Paste Raw Terminal Logs / Scan Output Here", 
    height=450, 
    placeholder="Example: Starting Nmap 7.92... \nNmap scan report for 192.168.1.1\nPORT 80/tcp open http\nPORT 443/tcp open https...")

with col2: 
    st.subheader("Report Generation") 
    st.write("Clicking the button below will analyze the logs for vulnerabilites, categorize them by risk, and draft a formal word document")

    generate_btn = st.button("Generate Professional Report")


# --- CORE LOGIC --- 
if generate_btn:
    if not raw_logs:
        st.error("Please provide scan data first.")
    else: 
        with st.spinner("AI is analyzing vulnerabilities and drafting export."):
            try:
                system_prompt = f"""You are a world-class Cybersecurity Consultant. Your task is to analyze raw technical logs and write a professional Pentest report for {client_name}, The Report must include: 1. An Executive Summary for the authorities (non-technical) 2. A Technical Findings for IT Staff. 3. Top 3 Vulnerabilities categorized as CRITICAL, HIGH or MEDIUM (if any). 4. Remediation steps for each. Keep the tone professional, authoritative, and clear."""
                chat_completion = client.chat.completions.create(
                    messages=[
                        {
                            "role": "system",
                            "content": system_prompt
                        },
                        {
                            "role": "user",
                            "content": f"Logs to analyze:\n{raw_logs}"
                        }
                    ], 
                    model=model_choice, 
                )
                report_content = chat_completion.choices[0].message.content

                from docx import Document
            
                doc = Document()

                doc.add_heading(f'Security Asessment: {client_name}', 0)
                doc.add_heading('Document Information', level=1)
                doc.add_paragraph(f"Date: {datetime.date.today().strftime('%B %d, %Y')}")

                doc.add_paragraph(f"Prepared By: {analyst_name}") 
                doc.add_paragraph(f"Audit Type: {report_type}") 

                doc.add_page_break()

                doc.add_heading('Executive Findings', level=1) 

                doc.add_paragraph(report_content)

                b_io = io.BytesIO()
                doc.save(b_io)

                st.success("Analysis Complete!")
                st.markdown("### Report Preview") 

                st.info(report_content[:1000] + "...") 

                st.download_button(
                    label="DOWNLOAD WORD REPORT (.docx)",
                    data = b_io.getvalue(),
                    file_name=f"Report_{client_name}_{datetime.date.today()}.docx",
                    mime="application/vnd.openxmlfromats-officedocument.wordprocessingml.document" 
                )
            except Exception as e:
                st.error(f"Error connecting to AI: {str(e)}")

# --- FOOTER --- 
st.markdown("---")
st.caption("Valedion AI - Proprietary Internal Security Tool. Data processed via secure groq LPU.")

